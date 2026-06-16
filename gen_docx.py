#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

with open('/home/asus/kuafu/九世元灵_第一卷.md', 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

lines = content.split('\n')
for line in lines:
    s = line.strip()
    if not s:
        continue
    if s.startswith('## '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s.replace('## ', ''))
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        continue
    if s.startswith('# '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s.replace('# ', ''))
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        continue
    if s == '---':
        continue
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    parts = s.split('**')
    for j, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.font.size = Pt(12)
        if j % 2 == 1:
            r.bold = True

output_path = '/home/asus/kuafu/九世元灵_第一卷.docx'
doc.save(output_path)
print(f'OK: {os.path.getsize(output_path)} bytes')
